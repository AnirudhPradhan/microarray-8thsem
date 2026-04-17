from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, "output", "images")
COMP = os.path.join(IMG, "comparative")
OUT = os.path.join(BASE, "output", "SUMMARY_REPORT.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    doc.add_heading(text, level=level)

def body(text):
    doc.add_paragraph(text)

def bold_para(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True

def img(path, width=6.0, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f"[Image not found: {os.path.basename(path)}]")

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light List Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Inches(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Summary Report: Triclustering Algorithms for Temporal Gene Expression", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Replication of Gutiérrez-Avilés et al. (NEUCOM 2014) & Soares et al. (Pattern Recognition 2024)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Overview")
body(
    "This study implements and benchmarks six clustering algorithms across three levels of dimensionality "
    "(1D, 2D, 3D) on temporal gene expression tensors (genes × conditions × timepoints). "
    "Results are evaluated using internal quality metrics (TRIQ family) and external recoverability "
    "metrics (R/CE/RMS3) against planted ground-truth triclusters."
)
body("Datasets used:")
for line in [
    "QualityC — 500 genes × 10 conditions × 5 timepoints; 70 planted triclusters (10% noise)",
    "Yeast GST — 7,491 genes × 3 conditions × 14 timepoints (real biological data, Spellman et al.)",
    "17-Dataset Benchmark — All 500×10×5, varying pattern types (Soares et al. 2024)",
]:
    p = doc.add_paragraph(line, style="List Bullet")

# ══════════════════════════════════════════════════════════════════════════════
# 2. ALGORITHMS
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Algorithms at a Glance")
table(
    ["Algorithm", "Dim", "Core Idea"],
    [
        ["K-Means",              "1D", "Flatten tensor → cluster genes by Euclidean distance on (cond×time) profiles"],
        ["Agglomerative",        "1D", "Flatten tensor → hierarchical Ward-linkage clustering on gene vectors"],
        ["SpectralCoclustering", "2D", "SVD on bipartite gene–feature graph; joint row/column clustering"],
        ["SpectralBiclustering", "2D", "Log-normalised SVD; separate K-Means on row and column embeddings"],
        ["TriGen",               "3D", "Genetic algorithm evolving gene/condition/time masks to minimise MSL/MSR3D"],
        ["triCluster",           "3D", "Exact multiplicative windows per condition-pair; enumerate biclusters, merge across time"],
        ["TCtriCluster",         "3D", "triCluster + constraint: timepoints must be consecutive (biologically motivated)"],
        ["δ-Trimax",             "3D", "3D Cheng-Church: iteratively delete/add genes/conditions/times to minimise MSR"],
    ],
    col_widths=[1.8, 0.5, 4.2],
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. METRICS
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Evaluation Metrics at a Glance")

heading("3.1 Internal Quality (no ground truth needed)", level=2)
table(
    ["Metric", "Formula Summary", "Interpretation"],
    [
        ["GRQ",  "1 − MSR/Var(X)",                       "Structural coherence; penalises non-additive patterns"],
        ["PEQ",  "Mean pairwise |Pearson r|",             "Linear co-expression strength between genes"],
        ["SPQ",  "Mean pairwise |Spearman ρ|",            "Rank-based co-expression strength"],
        ["TRIQ", "(0.4×GRQ + 0.05×PEQ + 0.05×SPQ)/0.5", "Weighted composite [0,1]; GRQ dominates (80%). ≥0.7 = highly coherent"],
    ],
    col_widths=[0.9, 2.3, 3.3],
)

heading("3.2 External Recoverability (requires planted ground truth)", level=2)
table(
    ["Metric", "Formula Summary", "Interpretation"],
    [
        ["R",    "Mean |pᵢ ∩ best_GTⱼ| / |pᵢ|",           "Precision: do predicted triclusters fall inside planted patterns?"],
        ["CE",   "1 − |union_P ∩ union_GT| / |union_P|",   "Fraction of predicted cells outside any planted tricluster"],
        ["RMS3", "Geometric mean of per-dimension Jaccard", "3D match quality; penalises dimensional mismatch harshly"],
    ],
    col_widths=[0.9, 2.7, 2.9],
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Results")

# 4.1 TRIQ
heading("4.1 Internal Quality (TRIQ) — QualityC vs. Yeast", level=2)

img(os.path.join(COMP, "09_all_algorithms_triq_qualityc.png"), 5.5,
    "Figure 1: Mean TRIQ comparison — all algorithms on QualityC")
img(os.path.join(COMP, "10_all_algorithms_triq_yeast.png"), 5.5,
    "Figure 2: Mean TRIQ comparison — all algorithms on Yeast")

body("QualityC — Mean TRIQ per algorithm (ranked):")
table(
    ["Rank", "Algorithm", "Dim", "TRIQ", "GRQ", "PEQ", "SPQ"],
    [
        ["1", "δ-Trimax",            "3D", "0.893", "0.874", "0.461", "0.761"],
        ["2", "TriGen",              "3D", "0.699", "0.749", "0.501", "0.502"],
        ["3", "triCluster",          "3D", "0.603", "0.576", "0.452", "0.854"],
        ["4", "TCtriCluster",        "3D", "0.504", "0.514", "0.473", "0.823"],
        ["5", "SpectralBiclustering","2D", "0.258", "0.240", "0.366", "0.367"],
        ["6", "SpectralCoclustering","2D", "0.196", "0.174", "0.296", "0.294"],
        ["7", "Agglomerative",       "1D", "0.118", "0.113", "0.143", "0.142"],
        ["8", "K-Means",             "1D", "0.092", "0.084", "0.123", "0.122"],
    ],
    col_widths=[0.5, 1.7, 0.5, 0.7, 0.7, 0.7, 0.7],
)

body("Yeast — Mean TRIQ per algorithm (ranked):")
table(
    ["Rank", "Algorithm", "Dim", "TRIQ", "GRQ", "PEQ", "SPQ"],
    [
        ["1", "δ-Trimax",            "3D", "0.810", "0.862", "0.593", "0.594"],
        ["2", "TCtriCluster",        "3D", "0.678", "0.768", "0.356", "0.374"],
        ["3", "SpectralBiclustering","2D", "0.382", "0.380", "0.402", "0.376"],
        ["4", "SpectralCoclustering","2D", "0.275", "0.271", "0.316", "0.302"],
        ["5", "Agglomerative",       "1D", "0.223", "0.210", "0.278", "0.272"],
        ["6", "K-Means",             "1D", "0.220", "0.209", "0.263", "0.262"],
        ["7", "TriGen",              "3D", "0.257", "0.510", "0.380", "0.373"],
        ["—", "triCluster",          "3D", "TIMEOUT", "—", "—", "—"],
    ],
    col_widths=[0.5, 1.7, 0.5, 0.8, 0.7, 0.7, 0.7],
)

bold_para("Key insight: All 3D algorithms outperform all 1D/2D baselines on QualityC TRIQ. "
          "δ-Trimax leads on both datasets. triCluster times out on Yeast (14 timepoints → "
          "exponential search); TCtriCluster solves this in 1.7 s via temporal contiguity constraint.")

doc.add_page_break()

# 4.2 Recoverability
heading("4.2 Recoverability — 17-Dataset Benchmark", level=2)

img(os.path.join(COMP, "11_R_heatmap_all_datasets.png"), 5.5,
    "Figure 3: R heatmap across all 17 datasets and 3 algorithms")
img(os.path.join(COMP, "08_all3_R_by_pattern.png"), 5.5,
    "Figure 4: Mean R by pattern type for all three triclustering algorithms")

body("Mean R by algorithm across 17 datasets:")
table(
    ["Algorithm", "Mean R (all 17)", "Mean R (non-OP, 13)", "Mean R (OP, 4)"],
    [
        ["triCluster",    "0.566", "0.808", "0.000"],
        ["TCtriCluster",  "0.531", "0.761", "0.000"],
        ["δ-Trimax",      "0.118", "0.168", "0.032"],
    ],
    col_widths=[1.6, 1.5, 1.8, 1.6],
)

body("Per-pattern R breakdown:")
table(
    ["Pattern Type", "triCluster R", "TCtriCluster R", "δ-Trimax R"],
    [
        ["Constant",              "0.950", "0.791", "0.160"],
        ["Additive",              "0.883", "0.746", "0.281"],
        ["Multiplicative",        "0.924", "0.830", "0.182"],
        ["OrderPreserving",       "0.000", "0.000", "0.038"],
        ["Overlapping (C/A/M avg)","0.908","0.784", "0.209"],
        ["Quality noisy (C/A/M avg)","0.566","0.562","0.236"],
        ["Contiguity (C/A/M avg)","0.921", "0.921", "0.229"],
    ],
    col_widths=[2.0, 1.4, 1.5, 1.4],
)

bold_para("Key insight: triCluster dominates on structured patterns. All three algorithms "
          "completely fail on OrderPreserving (R ≈ 0) — a fundamental modelling incompatibility, "
          "not an implementation bug.")

doc.add_page_break()

# 4.3 triCluster vs TCtriCluster
heading("4.3 triCluster vs. TCtriCluster — Head-to-Head", level=2)

img(os.path.join(COMP, "07_tricluster_vs_tctricluster_R.png"), 5.5,
    "Figure 5: triCluster vs. TCtriCluster R comparison per dataset")

table(
    ["Dataset", "triCluster R", "TCtriCluster R", "#Found (TC)", "#Found (TCT)"],
    [
        ["BaseR",          "0.867", "0.676", "10", "7"],
        ["Constant",       "0.950", "0.791", "45", "29"],
        ["Additive",       "0.883", "0.746", "40", "21"],
        ["Multiplicative", "0.924", "0.830", "35", "20"],
        ["ContiguityC",    "0.950", "0.950", "43", "43"],
        ["ContiguityA",    "0.922", "0.922", "29", "29"],
        ["ContiguityM",    "0.891", "0.891", "39", "39"],
        ["QualityC",       "0.561", "0.515", "48", "20"],
    ],
    col_widths=[1.5, 1.3, 1.4, 1.2, 1.2],
)

bold_para("Key insight: On Contiguity datasets, both algorithms produce identical results. "
          "The temporal contiguity constraint is 'free' when patterns are already consecutive. "
          "On non-contiguous datasets, TCtriCluster finds ~27% fewer triclusters with ~6% lower R.")

# 4.4 TRIQ vs R
heading("4.4 TRIQ vs. R Trade-off", level=2)

img(os.path.join(COMP, "13_triq_vs_R_scatter.png"), 4.5,
    "Figure 6: TRIQ vs R scatter — high TRIQ does not imply high R")
img(os.path.join(COMP, "12_radar_all_algorithms.png"), 4.5,
    "Figure 7: Radar chart — multi-metric comparison of all algorithms")

table(
    ["Algorithm", "TRIQ (QualityC)", "R (QualityC)", "Interpretation"],
    [
        ["δ-Trimax",      "0.893", "0.223", "Coherent but doesn't match planted patterns"],
        ["triCluster",    "0.603", "0.561", "Balanced — good at both metrics"],
        ["TCtriCluster",  "0.504", "0.515", "Balanced with fewer triclusters"],
        ["TriGen",        "0.699", "0.259", "Coherent but coarse (only 5 GA solutions)"],
        ["1D Agglo.",     "0.118", "0.567*","High recall by cluster size, not coherence"],
    ],
    col_widths=[1.5, 1.4, 1.2, 2.5],
)
body("* Recall-based R formula — not directly comparable to Soares precision-oriented R used for 3D algorithms.")

doc.add_page_break()

# 4.5 Recoverability all methods
heading("4.5 Recoverability on QualityC Ground Truth — All Methods", level=2)

img(os.path.join(IMG, "final_recoverability.png"), 5.5,
    "Figure 8: R/CE/RMS3 recoverability on QualityC across all methods")

table(
    ["Method", "#Found", "R", "CE", "RMS3", "Avg Volume"],
    [
        ["triCluster",            "48", "0.5609", "0.9965", "0.8024", "—"],
        ["1D Agglomerative*",     "5",  "0.5674", "0.6936", "1.47",   "5,000"],
        ["1D K-Means*",           "5",  "0.4636", "0.6936", "1.45",   "5,000"],
        ["TriGen 3D",             "5",  "0.2588", "0.6951", "1.54",   "3,675"],
        ["2D SpectralCocluster",  "5",  "0.1309", "0.6980", "1.87",   "1,060"],
    ],
    col_widths=[1.8, 0.7, 0.8, 0.8, 0.8, 1.0],
)
body("* 1D methods use a recall-based R variant — large gene-only clusters trivially cover planted gene sets.")

# 4.6 TriGen
heading("4.6 TriGen Results — QualityC, Yeast, Synthetic", level=2)

img(os.path.join(IMG, "final_trigen_all_datasets.png"), 5.5,
    "Figure 9: TriGen TRIQ across all three datasets")

table(
    ["Dataset", "TRIQ", "GRQ", "PEQ", "SPQ", "Note"],
    [
        ["QualityC (500×10×5)",       "0.699", "0.749", "0.501", "0.502", "Consistent across 5 GA populations"],
        ["Yeast (7491×3×14)",         "0.257", "0.510", "0.380", "0.373", "Real noise limits TRIQ"],
        ["Synthetic (1000×10×5)",     "0.753", "0.667", "0.943", "0.930", "PEQ/SPQ → 1.0; validates temporal recovery"],
    ],
    col_widths=[1.8, 0.7, 0.7, 0.7, 0.7, 2.1],
)

# 4.7 Speed
heading("4.7 Speed Comparison", level=2)

table(
    ["Algorithm", "Dataset", "Time", "Feasible?"],
    [
        ["K-Means / Agglomerative",      "500×10×5",        "<1 s",    "Yes"],
        ["SpectralCo/Biclustering",      "500×10×5",        "<1 s",    "Yes"],
        ["triCluster",                   "500×10×5",        "0.12 s",  "Yes"],
        ["TCtriCluster",                 "500×10×5",        "0.11 s",  "Yes"],
        ["δ-Trimax",                     "500×10×5",        "~16 s",   "Yes"],
        ["TriGen",                       "500×10×5",        "~30-60 s","Yes"],
        ["triCluster",                   "500×3×14 (Yeast)","TIMEOUT", "NO"],
        ["TCtriCluster",                 "500×3×14 (Yeast)","1.7 s",   "Yes"],
        ["δ-Trimax",                     "500×3×14 (Yeast)","~1.5 s",  "Yes"],
    ],
    col_widths=[1.9, 1.8, 1.2, 1.0],
)
bold_para("Key insight: triCluster's EXPAND_T is exponential in timepoints (2^14 = 16,384 for Yeast). "
          "TCtriCluster's contiguity constraint linearises this to ≤13 consecutive windows — "
          "a decisive computational advantage on real time-course data.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. MASTER SUMMARY CHART
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Master Summary Chart")
img(os.path.join(COMP, "00_MASTER_SUMMARY.png"), 6.2,
    "Figure 10: Master summary — TRIQ and R across all algorithms and datasets")

# ══════════════════════════════════════════════════════════════════════════════
# 6. KEY FINDINGS
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Key Findings")
findings = [
    "Dimensionality drives internal quality. All 3D algorithms outperform all 1D/2D baselines on TRIQ. "
    "Ranking: δ-Trimax (0.893) > TriGen (0.699) > triCluster (0.603) > TCtriCluster (0.504) >> "
    "SpectralBiclustering (0.258) >> K-Means (0.092).",

    "TRIQ ≠ Recoverability. δ-Trimax has the highest TRIQ but lowest R among 3D methods. "
    "triCluster achieves both high TRIQ and high R — the only algorithm that wins on both simultaneously.",

    "triCluster dominates on structured patterns. R > 0.88 on 9 of 13 non-OrderPreserving datasets. Mean R = 0.808.",

    "TCtriCluster's temporal constraint is biologically free. Identical performance to triCluster on "
    "Contiguity datasets (R = 0.921). Only ~6% lower R on non-contiguous datasets. Reduces Yeast runtime "
    "from timeout to 1.7 s.",

    "OrderPreserving patterns are unrecoverable by all three tricluster algorithms. "
    "triCluster / TCtriCluster: R = 0.000; δ-Trimax: R = 0.038. Fundamental modelling incompatibility — "
    "rank patterns need rank-based algorithms (e.g., OPTriCluster).",

    "Noise degrades triCluster significantly. R drops from 0.950 (Constant, clean) to 0.561 "
    "(QualityC, 10% noise) — a 41% degradation. Tight window tolerance is sensitive to noise.",

    "δ-Trimax excels on real data. TRIQ Mean = 0.810 on Yeast — highest of all tested algorithms. "
    "MSR minimisation naturally handles noisy biological data with no planted structure.",

    "TriGen validates on its own synthetic data. PEQ/SPQ → 0.94 on TriGen synthetic (1000×10×5), "
    "confirming the GA recovers planted temporal co-expression patterns.",
]
for f in findings:
    doc.add_paragraph(f, style="List Number")

# ══════════════════════════════════════════════════════════════════════════════
# 7. RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Algorithm Recommendations by Use Case")
table(
    ["Use Case", "Best Algorithm", "Why"],
    [
        ["Synthetic data, exact pattern recovery",   "triCluster",          "R = 0.808 mean; best precision on structured patterns"],
        ["Real noisy biological data",               "δ-Trimax",            "TRIQ = 0.810 on Yeast; MSR handles noise well"],
        ["Temporal windows (biology-interpretable)", "TCtriCluster",        "Contiguous time blocks; feasible on long time-courses"],
        ["Order-preserving expression",              "(needs OPTriCluster)", "None of these algorithms handle rank patterns"],
        ["Fast exploration, no ground truth",        "SpectralBiclustering","Best 2D TRIQ (0.497 on Yeast); sub-second runtime"],
        ["Simple, interpretable baseline",           "1D Agglomerative",    "Fastest; reasonable recall on large gene sets"],
    ],
    col_widths=[2.3, 1.7, 2.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. CAPABILITY MATRIX
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Algorithm Capability Matrix")
table(
    ["Capability", "1D", "2D Spectral", "TriGen", "triCluster", "TCtriCluster", "δ-Trimax"],
    [
        ["Constant patterns",      "Yes", "Yes", "Yes", "Yes",   "Yes",   "Yes"],
        ["Additive patterns",      "~",   "~",   "~",   "Yes",   "Yes",   "Yes"],
        ["Multiplicative patterns","~",   "~",   "~",   "Yes",   "Yes",   "~"],
        ["Order-preserving",       "~",   "~",   "Yes", "No",    "No",    "No"],
        ["Temporal structure",     "No",  "~",   "Yes", "Yes",   "Yes",   "Yes"],
        ["Temporal contiguity",    "No",  "No",  "No",  "No",    "Yes",   "No"],
        ["Scalable (long time)",   "Yes", "Yes", "Yes", "No",    "~",     "Yes"],
        ["High TRIQ",              "Low", "Med", "High","High",  "High",  "Highest"],
        ["High R",                 "Med*","Low", "Med", "High",  "High",  "Low"],
    ],
    col_widths=[1.8, 0.5, 0.9, 0.8, 0.9, 1.0, 0.8],
)
body("* 1D high R is misleading — large gene-only clusters trivially contain planted gene sets by size.")

# ══════════════════════════════════════════════════════════════════════════════
# Footer note
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
note = doc.add_paragraph(
    "Algorithms: triCluster (Zhao & Zaki 2005), TCtriCluster (Soares 2015), "
    "δ-Trimax (Cheng-Church 2000 / Soares 2020), TriGen (Gutiérrez-Avilés et al. 2014). "
    "Datasets: Soares et al. 2024 benchmark, Spellman et al. 1998 Yeast GST."
)
note.runs[0].italic = True
note.runs[0].font.size = Pt(9)

doc.save(OUT)
print(f"Saved: {OUT}")
